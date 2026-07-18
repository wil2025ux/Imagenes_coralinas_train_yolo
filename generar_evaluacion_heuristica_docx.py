#!/usr/bin/env python3
"""Genera documento Word con evaluación heurística Nielsen del prototipo CoArrCP IA."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent / "ENTREGA_TESIS_SEGMENTACION_BENTONICA" / "documentos" / "Evaluacion_Heuristica_CoArrCP_IA.docx"


def set_cell_shading(cell, hex_color: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row):
            cells[c_idx].text = str(text)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    return table


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("Evaluación heurística de usabilidad", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Prototipo: CoArrCP IA Flask — Segmentación bentónica con YOLO-seg")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph(
        "Proyecto de tesis: segmentación automatizada de componentes bentónicos "
        "(algas, corales, almejas, esponjas, arena) en imágenes coralinas, "
        "con flujo inspirado en CoArrCP."
    )

    doc.add_heading("Metodología: 5 pasos", level=1)
    steps = [
        ("PASO 1 · Evaluación 1", "Evaluar el prototipo ORIGINAL (mini Flask simple: subir imagen, ver resultado, sin menú CoArrCP ni puntos editables)."),
        ("PASO 2 · Resultados", "Identificar heurísticas con calificación 1 (en desacuerdo) como prioridad de mejora."),
        ("PASO 3 · Cambios en el prototipo", "Aplicar mejoras según recomendaciones de Nielsen y documentarlas."),
        ("PASO 4 · Evaluación 2", "Evaluar el prototipo MEJORADO (versión actual con menú lateral, SQLite, drag-and-drop y puntos editables)."),
        ("PASO 5 · Comparar resultados", "Analizar diferencias entre Evaluación 1 y Evaluación 2."),
    ]
    for name, desc in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    doc.add_heading("Escala Likert de 3 puntos", level=1)
    likert = doc.add_table(rows=4, cols=2)
    likert.style = "Table Grid"
    likert.rows[0].cells[0].text = "Calificación"
    likert.rows[0].cells[1].text = "Significado"
    likert.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    likert.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    likert_data = [
        ("1 — En desacuerdo", "El prototipo NO cumple el criterio. Se requieren mejoras importantes.", "FFCCCC"),
        ("2 — Neutral", "El prototipo cumple PARCIALMENTE. Hay oportunidades de mejora.", "FFF2CC"),
        ("3 — De acuerdo", "El prototipo cumple SATISFACTORIAMENTE. No requiere cambios inmediatos.", "D9EAD3"),
    ]
    for i, (score, meaning, color) in enumerate(likert_data, start=1):
        likert.rows[i].cells[0].text = score
        likert.rows[i].cells[1].text = meaning
        set_cell_shading(likert.rows[i].cells[0], color)

    doc.add_heading("Evaluación 1 — Prototipo ORIGINAL", level=1)
    doc.add_paragraph(
        "Escala: 1 = En desacuerdo · 2 = Neutral · 3 = De acuerdo. "
        "Promedio aproximado: 1.6"
    )

    eval1_headers = ["#", "Heurística (Nielsen)", "Nota", "Justificación"]
    eval1_rows = [
        ("1", "Visibilidad del estado del sistema", "2", "Mostraba resultados de IA, pero no estado claro del modelo, progreso al procesar ni flujo general."),
        ("2", "Correspondencia con el mundo real", "2", "Clases bentónicas correctas, pero términos técnicos (conf, iou) sin explicación en pantalla."),
        ("3", "Control y libertad del usuario", "1", "No se podían mover puntos, ni marcar revisado, ni corregir inferencias manualmente."),
        ("4", "Consistencia y estándares", "2", "Interfaz funcional pero distinta al flujo CoArrCP esperado por el usuario objetivo."),
        ("5", "Prevención de errores", "1", "Había que ejecutar comandos en terminal; fácil confundirse si el servidor no estaba activo."),
        ("6", "Reconocimiento antes que recuerdo", "2", "Resultados visibles en imagen, pero sin menú ni recordatorio de pasos del flujo."),
        ("7", "Flexibilidad y eficiencia de uso", "1", "Solo carga archivo a archivo; sin carpeta, sin arrastrar, sin revisión por lotes."),
        ("8", "Diseño estético y minimalista", "2", "Página simple, poca jerarquía visual y sin identidad tipo CoArrCP."),
        ("9", "Ayuda a reconocer, diagnosticar y corregir errores", "1", "Errores poco claros (modelo no cargado, servidor caído) sin guía de solución."),
        ("10", "Ayuda y documentación", "1", "Sin ayuda integrada sobre parámetros, puntos de conteo ni uso del sistema."),
    ]
    t1 = add_table(doc, eval1_headers, eval1_rows, [0.4, 2.2, 0.5, 3.4])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Prioridades identificadas (Paso 2): ").bold = True
    p.add_run("Control del usuario, prevención de errores, ayuda/documentación y flexibilidad de uso.")

    doc.add_heading("Paso 3 — Cambios realizados", level=1)
    cambios = [
        "Menú lateral tipo CoArrCP (inicio, carga, inferencia, revisión, componentes, base de datos, estadística, configuración).",
        "Carga por arrastre (drag-and-drop) de imágenes y carpetas completas.",
        "Puntos de conteo visibles (verde = detectado, rojo = sin detección) y editables arrastrando sobre componentes bentónicos.",
        "Persistencia en SQLite e historial de inferencias, coberturas, detecciones y puntos.",
        "Flujo de revisión manual con estados pendiente / revisado.",
        "Un solo comando de arranque (run.py) y modo --reload para desarrollo sin reiniciar manualmente.",
        "Leyenda de colores y tabla de puntos sincronizada al mover cada punto (clase recalculada automáticamente).",
        "Mapeo de clases YOLO a componentes CoArrCP (GGMF, GMF) en pestaña Componentes.",
    ]
    for c in cambios:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("Evaluación 2 — Prototipo MEJORADO", level=1)
    doc.add_paragraph(
        "Escala: 1 = En desacuerdo · 2 = Neutral · 3 = De acuerdo. "
        "Promedio aproximado: 2.7"
    )

    eval2_headers = ["#", "Heurística (Nielsen)", "Nota", "Justificación"]
    eval2_rows = [
        ("1", "Visibilidad del estado del sistema", "3", "Indicador Modelo cargado, métricas en inicio, mensajes flash y estado al guardar puntos."),
        ("2", "Correspondencia con el mundo real", "2", "Flujo alineado a CoArrCP; conf/iou siguen siendo técnicos (falta tooltip o ayuda contextual)."),
        ("3", "Control y libertad del usuario", "3", "Arrastrar puntos, marcar revisado/pendiente, limpiar selección, elegir carpeta o archivos."),
        ("4", "Consistencia y estándares", "3", "Menú y etapas coherentes con el flujo de trabajo bentónico del proyecto."),
        ("5", "Prevención de errores", "2", "Validación al subir sin archivos; aún requiere terminal para arrancar (no es app empaquetada)."),
        ("6", "Reconocimiento antes que recuerdo", "3", "Menú siempre visible, leyenda de colores, tabla de puntos vinculada a la imagen."),
        ("7", "Flexibilidad y eficiencia de uso", "3", "Carpeta completa, drag-and-drop, procesamiento por lote, estadísticas globales."),
        ("8", "Diseño estético y minimalista", "3", "UI oscura organizada, tarjetas, métricas y panel de revisión claro."),
        ("9", "Ayuda a reconocer, diagnosticar y corregir errores", "2", "Mensajes flash útiles; faltan textos guiados si YOLO no carga o falla inferencia."),
        ("10", "Ayuda y documentación", "2", "README y API documentados; no hay ayuda contextual dentro de la app."),
    ]
    add_table(doc, eval2_headers, eval2_rows, [0.4, 2.2, 0.5, 3.4])

    doc.add_heading("Paso 5 — Comparación de resultados", level=1)
    doc.add_paragraph(
        "Entre Evaluación 1 y Evaluación 2 mejoraron sobre todo: control del usuario (1→3), "
        "flexibilidad y eficiencia (1→3), visibilidad del estado (2→3) y diseño estético (2→3). "
        "Permanecen en calificación 2 las heurísticas ligadas al arranque técnico (terminal), "
        "parámetros sin explicar en la interfaz y ayuda integrada."
    )

    doc.add_heading("Comparativa resumida (Evaluación 1 → Evaluación 2)", level=2)
    cmp_headers = ["Heurística", "Eval. 1", "Eval. 2", "Variación"]
    cmp_rows = [
        ("Visibilidad del estado", "2", "3", "+1"),
        ("Correspondencia mundo real", "2", "2", "0"),
        ("Control y libertad", "1", "3", "+2"),
        ("Consistencia y estándares", "2", "3", "+1"),
        ("Prevención de errores", "1", "2", "+1"),
        ("Reconocimiento vs recuerdo", "2", "3", "+1"),
        ("Flexibilidad y eficiencia", "1", "3", "+2"),
        ("Diseño estético", "2", "3", "+1"),
        ("Diagnóstico y recuperación de errores", "1", "2", "+1"),
        ("Ayuda y documentación", "1", "2", "+1"),
        ("PROMEDIO", "1.6", "2.7", "+1.1"),
    ]
    add_table(doc, cmp_headers, cmp_rows, [2.8, 0.8, 0.8, 0.8])

    doc.add_heading("Reflexión final", level=1)
    doc.add_paragraph(
        "El prototipo pasó de ser un validador técnico de YOLO-seg a una mini réplica funcional "
        "del flujo CoArrCP, con carga operativa, revisión de puntos y trazabilidad en base de datos. "
        "Las heurísticas con nota 1 en la evaluación inicial reflejaban la brecha entre "
        "'funciona para el desarrollador' y 'es usable para el analista bentónico'; la versión "
        "mejorada reduce esa brecha, aunque aún no alcanza producto final (integración Xojo + "
        "despliegue sin terminal)."
    )

    doc.add_heading("Próximas mejoras sugeridas", level=1)
    for item in [
        "Tooltips en conf, iou y Puntos dentro de la pantalla Cargar imágenes.",
        "Sección de ayuda integrada en Configuración.",
        "Empaquetado o instalador para no depender de la terminal.",
        "Integración con Xojo como interfaz definitiva de producción.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    main()
