#!/usr/bin/env python3
"""Genera informe Word en voz de biólogo experto: interfaz y atajos CoArrCP IA."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUTPUT = (
    Path(__file__).resolve().parent
    / "ENTREGA_TESIS_SEGMENTACION_BENTONICA"
    / "documentos"
    / "Informe_Biologo_Interfaz_y_Atajos_CoArrCP_IA.docx"
)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.rows[r + 1].cells[c].text = str(text)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Portada
    title = doc.add_heading("Informe de uso del prototipo CoArrCP IA", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Documentación de interfaz y teclas rápidas")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Elaborado por: ").bold = True
    meta.add_run("Biólogo marino — usuario experto en conteo bentónico y revisión de cuadrantes coralinos\n")
    meta.add_run("Herramienta evaluada: ").bold = True
    meta.add_run("CoArrCP IA Flask (segmentación YOLO-seg de algas, corales, almejas, esponjas y arena)\n")
    meta.add_run("Fecha de evaluación: ").bold = True
    meta.add_run("Mayo 2026\n")
    meta.add_run("Acceso: ").bold = True
    meta.add_run("http://127.0.0.1:9001")

    doc.add_page_break()

    # Introducción
    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(
        "Como biólogo que revisa cientos de fotografías bentónicas por campaña, necesito una "
        "herramienta donde la imagen del fondo marino sea lo principal y no me distraigan menús "
        "técnicos. Evalué el prototipo CoArrCP IA durante una sesión de trabajo simulando mi rutina "
        "habitual: cargar transectos fotográficos, verificar la segmentación automática, corregir "
        "puntos de conteo sobre corales y algas, marcar cada foto como revisada y avanzar a la "
        "siguiente lo más rápido posible."
    )
    doc.add_paragraph(
        "Este informe documenta cómo uso la interfaz, qué significa cada pantalla para mi trabajo "
        "de campo y laboratorio, y cómo empleo las teclas rápidas para no depender del ratón en "
        "tareas repetitivas."
    )

    doc.add_heading("2. Mi flujo de trabajo diario", level=1)
    add_table(
        doc,
        ["Paso", "Qué hago yo", "Dónde en la app", "Atajo / gesto"],
        [
            ("1", "Llego con una carpeta de fotos del arrecife", "Cargar imágenes", "Arrastrar carpeta al recuadro"),
            ("2", "Ejecuto la IA sobre todo el lote", "Cargar imágenes", "Clic en «Procesar con IA»"),
            ("3", "Compruebo que se procesaron todas", "Inferencia IA", "—"),
            ("4", "Abro la primera foto para revisar", "Revisión manual", "Clic en «Revisar»"),
            ("5", "Coloco los puntos sobre el bentos real", "Revisión manual", "Arrastrar círculos numerados"),
            ("6", "Doy por buena la foto", "Revisión manual", "Tecla R"),
            ("7", "Paso a la siguiente sin tocar el ratón", "Revisión manual", "Tecla K o →"),
            ("8", "Guardo una copia para el informe", "Revisión manual", "Tecla D"),
        ],
        [0.5, 2.0, 1.5, 1.5],
    )
    doc.add_paragraph(
        "Secuencia que más uso cuando ya tengo ritmo: reviso puntos → R → K → R → K. "
        "Con decenas o cientos de imágenes, esto reduce mucho la fatiga comparado con ir "
        "clicando menús y botones en cada foto."
    ).runs[0].italic = True

    doc.add_heading("3. Descripción de la interfaz (mi perspectiva)", level=1)

    sections = [
        (
            "3.1 Menú lateral",
            "A la izquierda veo las etapas del trabajo. Para mí el orden natural es: "
            "Cargar imágenes → Inferencia IA → Revisión manual. Las opciones Componentes, "
            "Base de datos, Estadística y Configuración las uso al inicio del proyecto o "
            "al cerrar una campaña, no en cada foto. El indicador «Modelo cargado» me confirma "
            "que la IA está lista antes de procesar un lote.",
        ),
        (
            "3.2 Inicio",
            "Es mi tablero de control. Veo cuántas imágenes llevo, cuántas faltan por revisar "
            "y cuántas ya validé. Desde aquí salto directo a cargar más fotos o retomo la última "
            "pendiente. No me quedo mucho tiempo aquí; es un resumen útil al abrir la sesión.",
        ),
        (
            "3.3 Cargar imágenes",
            "Esta es la pantalla de entrada de campaña. Lo que más valoro es poder arrastrar "
            "toda la carpeta del día de muestreo. Los parámetros conf, iou y número de puntos "
            "están ocultos en «Parámetros avanzados» — bien, porque yo no los toco en cada "
            "corrida; el técnico ya los dejó calibrados. Así la pantalla se centra en lo "
            "importante: seleccionar fotos y procesar.",
        ),
        (
            "3.4 Inferencia IA",
            "Aquí verifico que el lote terminó. Reviso cuántas detecciones hubo por imagen y "
            "el tiempo de procesamiento. Si algo falló, lo noto antes de sentarme a revisar "
            "punto por punto. El botón «Revisar» me lleva directo a cada foto.",
        ),
        (
            "3.5 Revisión manual (pantalla principal de mi trabajo)",
            "Esta es la pantalla crítica. Al abrir una foto entra el modo foco: la fotografía "
            "del fondo marino ocupa casi toda la pantalla y desaparece el menú lateral. "
            "Encima veo círculos numerados (13 puntos de conteo): verdes si la IA detectó "
            "componente en ese sitio, rojos si no. Arrastro cada punto hasta donde yo sé que "
            "hay coral, alga, almeja, esponja o arena; al soltar, la clase se recalcula sola.\n\n"
            "Arriba hay una barra fina con: contador de foto (ej. 3/24), nombre del archivo, "
            "estado (pendiente/revisado) y botones equivalentes a las teclas rápidas. "
            "Los paneles de cobertura y lista de imágenes vienen cerrados; solo los abro "
            "si necesito comparar números (tecla G) o saltar a otra foto lejana (tecla L).",
        ),
        (
            "3.6 Componentes",
            "Relaciono las clases que devuelve la IA con los nombres que uso en CoArrCP "
            "(GGMF, GMF). Lo configuro una vez al inicio del proyecto bentónico, no en cada sesión.",
        ),
        (
            "3.7 Estadística y Base de datos",
            "Al terminar una campaña reviso coberturas promedio y frecuencia de clases. "
            "Me sirve para el informe final. La base SQLite queda en el equipo; no necesito "
            "interactuar con ella en el día a día.",
        ),
    ]
    for heading, text in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)

    doc.add_heading("4. Teclas rápidas que uso como biólogo experto", level=1)
    doc.add_paragraph(
        "Aprendí estas teclas en la primera media hora. Después de eso, en revisión casi "
        "no uso el ratón excepto para arrastrar puntos sobre el bentos."
    )

    add_table(
        doc,
        ["Tecla", "Acción", "Cuándo la uso yo"],
        [
            ("J o ←", "Foto anterior", "Cuando me equivoqué de imagen o quiero comparar con la previa"),
            ("K o →", "Foto siguiente", "Tras marcar revisada; es mi tecla más usada"),
            ("1 – 9", "Resaltar punto N", "Para localizar rápido un punto en la tabla o en la foto"),
            ("R", "Marcar revisada", "Cuando los puntos coinciden con lo que veo en el transecto"),
            ("U", "Marcar pendiente", "Si dudo y quiero volver después (mala luz, sedimento, etc.)"),
            ("D", "Descargar JPG etiquetada", "Para adjuntar al informe o compartir con el equipo"),
            ("G", "Mostrar/ocultar datos", "Solo cuando necesito el % de cobertura por clase"),
            ("L", "Mostrar/ocultar lista", "Para saltar a una foto específica del lote"),
            ("F", "Modo foco", "Alternar menú lateral; en revisión ya entra solo en foco"),
            ("?", "Ayuda de atajos", "La primera vez; luego ya no la necesito"),
        ],
        [0.9, 1.8, 3.8],
    )

    doc.add_heading("5. Significado de los puntos de conteo", level=1)
    doc.add_paragraph(
        "Los 13 puntos imitan el muestreo por puntos que conozco de metodologías bentónicas. "
        "Cada círculo numerado representa un punto de conteo sobre la imagen:"
    )
    add_table(
        doc,
        ["Color", "Significado", "Qué hago yo"],
        [
            ("Verde", "Componente detectado", "Verifico si la IA acertó; si no, arrastro el punto"),
            ("Rojo", "Sin detección", "Muevo el punto hasta un coral, alga u otro componente visible"),
        ],
        [1.0, 2.0, 3.5],
    )
    doc.add_paragraph(
        "Las cinco clases que reconoce el modelo son: algas, corales, almejas, esponjas y arena. "
        "Al mover un punto, la clase se actualiza según la máscara de segmentación en esa posición."
    )

    doc.add_heading("6. Clases bentónicas del modelo", level=1)
    add_table(
        doc,
        ["Clase IA", "Componente típico en mi transecto"],
        [
            ("Algas", "Macroalgas, turf algal, algas calcáreas"),
            ("Corales", "Coral hermatípico, colonias visibles"),
            ("Almejas", "Bivalvos sobre el sustrato"),
            ("Esponjas", "Esponjas incrustantes o tubulares"),
            ("Arena", "Sustrato arenoso sin organismo"),
        ],
        [1.5, 4.0],
    )

    doc.add_heading("7. Mi opinión como usuario experto", level=1)

    doc.add_heading("Lo que funciona bien", level=2)
    for item in [
        "La foto del arrecife domina la pantalla en revisión; es lo que necesito ver horas seguidas.",
        "Ocultar conf/iou en carga evita errores y distracciones.",
        "K y R me permiten un ritmo continuo sin buscar botones.",
        "Los puntos verdes/rojos son intuitivos para validar rápido.",
        "Arrastrar puntos es natural: coloco el conteo donde yo lo haría a ojo.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Lo que mejoraría", level=2)
    for item in [
        "Atajo para puntos 10–13 (ahora 1–9 cubre los primeros nueve).",
        "Indicador más visible de «guardando» al soltar un punto en conexión lenta.",
        "Exportar lote completo revisado en un clic (hoy descargo foto por foto con D).",
        "Breve texto de ayuda sobre qué es conf/iou si algún colega abre Parámetros avanzados.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Valoración para trabajo repetitivo", level=2)
    add_table(
        doc,
        ["Aspecto", "Nota (1–5)", "Comentario del evaluador"],
        [
            ("Velocidad para revisar una foto", "4", "Con K y R el flujo es ágil; arrastrar puntos sigue siendo manual"),
            ("Claridad de la interfaz", "4", "Modo foco ayuda; la barra superior podría ser aún más discreta"),
            ("Utilidad de teclas rápidas", "5", "Esenciales para no cansarme en lotes grandes"),
            ("Confianza en la segmentación IA", "3–4", "Depende de la foto; por eso reviso siempre los puntos"),
            ("¿La usaría en campaña real?", "Sí", "Como apoyo a revisión, no sustituto del criterio del biólogo"),
        ],
        [2.2, 0.8, 3.5],
    )

    doc.add_heading("8. Conclusión", level=1)
    doc.add_paragraph(
        "Desde mi rol de biólogo experto, el prototipo CoArrCP IA demuestra que puede integrarse "
        "en un flujo de trabajo continuo de revisión bentónica. La interfaz prioriza correctamente "
        "la fotografía submarina sobre menús y parámetros técnicos. Las teclas J, K, R y D forman "
        "el núcleo de mi operación diaria; el ratón lo reservo casi solo para ajustar puntos de "
        "conteo sobre los componentes reales del fondo marino."
    )
    doc.add_paragraph(
        "Recomiendo validar con una sesión de campo piloto de al menos 50 imágenes antes de "
        "integrar la herramienta con el sistema definitivo (CoArrCP / Xojo). Este informe puede "
        "servir como documentación de la opinión del usuario experto solicitada en la metodología "
        "de evaluación heurística y diseño centrado en el usuario."
    )

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("\n\n_______________________________\n").bold = False
    sig.add_run("Firma del biólogo evaluador\n")
    sig.add_run("Nombre: _______________________\n")
    sig.add_run("Institución: ___________________")

    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    main()
