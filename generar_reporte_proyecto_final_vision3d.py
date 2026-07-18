#!/usr/bin/env python3
"""
Genera el Reporte Técnico del Proyecto Final:
Procesamiento y reconocimiento visual en 3D — Vía 1 (2D→asistente 3D/future)
CoArrCP IA + YOLO-seg + SAM3 (Colab).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
DOCS = ROOT / "ENTREGA_TESIS_SEGMENTACION_BENTONICA" / "documentos"
OUT = DOCS / "Reporte_Tecnico_Proyecto_Final_Vision_3D_CoArrCP_IA.docx"
MD_OUT = ROOT / "Reporte_Tecnico_Proyecto_Final_Vision_3D_CoArrCP_IA.md"
FIGS = DOCS / "figuras_reporte"


def add_figure(doc, path: Path, caption: str, width_in: float = 5.8):
    """Inserta imagen centrada + pie de figura. Omite si no existe."""
    path = Path(path)
    if not path.is_file():
        p(doc, f"[Figura no encontrada: {path.name}]", bold=True)
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def build():
    doc = Document()

    # ---------- PORTADA ----------
    h = doc.add_heading("Reporte Técnico — Proyecto Final", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Procesamiento y reconocimiento visual en 3D"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Segmentación bentónica asistida: CoArrCP IA, YOLO-seg y SAM3 (Vía 1)"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for label, val in [
        ("Autor:", "Wilder Arath Utrera Vargas"),
        ("Sistema:", "CoArrCP IA (Flask) + Ultralytics YOLO-seg + Meta SAM3"),
        ("Vía de extracción:", "Vía 1 — De 2D a 3D (segmentación avanzada 2D; reconstrucción 3D como trabajo futuro)"),
        ("Fecha:", "Julio 2026"),
    ]:
        para = doc.add_paragraph()
        para.add_run(label + " ").bold = True
        para.add_run(val)

    doc.add_heading("Cumplimiento frente a los lineamientos del curso", level=2)
    add_table(
        doc,
        ["Requisito del curso", "Cómo se cubre en este proyecto"],
        [
            (
                "Arreglo experimental (videos/imágenes)",
                "Captura submarina 2021-06-08 con Nikon AW130/W300 (EXIF); organización tipo CoArrCP; §3.2.1 documenta obtención y requisitos de buceo/equipo.",
            ),
            (
                "Procesamiento y segmentación — Vía 1",
                "SAM3 (prompt texto/punto) + YOLO11n-seg entrenado en clases bentónicas; máscaras exportables YOLO-seg.",
            ),
            (
                "Detección y clasificación final",
                "YOLO-seg clasifica algas/corales/almejas/esponjas/arena; CoArrCP IA proyecta 13 puntos de conteo y permite corrección humana.",
            ),
            (
                "Manual del arreglo experimental",
                "Sección 3.4 de este reporte (hardware, entorno, software, procedimiento), adaptado de la Guía rápida CoArrCP (Xojo) al stack actual.",
            ),
            (
                "Reporte con estructura obligatoria",
                "Introducción, Trabajos relacionados, Materiales y métodos, Resultados, Conclusiones, Trabajo futuro.",
            ),
            (
                "Presentación oral",
                "No aplica en esta entrega (según acuerdo del autor); el reporte concentra metodología y evidencia.",
            ),
        ],
        widths=[2.2, 4.5],
    )

    # ---------- 1 INTRODUCCIÓN ----------
    doc.add_heading("1. Introducción", level=1)
    p(
        doc,
        "El monitoreo de coberturas bentónicas en arrecifes coralinos tradicionalmente se realiza mediante "
        "conteo de puntos sobre fotografías de fondo (método implementado por décadas en el software CoArrCP "
        "de Vargas Hernández). Ese flujo es científicamente sólido, pero intensivo en tiempo humano: cada "
        "imagen requiere identificar componentes en 13 (o más) puntos, con catálogos de más de 100 clases "
        "morfofuncionales.",
    )
    p(
        doc,
        "Este proyecto final aborda el ciclo de visión computacional espacial exigido por la experiencia "
        "educativa Procesamiento y reconocimiento visual en 3D, eligiendo la Vía 1 (De 2D a 3D): primero se "
        "aísla el objeto de interés en el dominio 2D con modelos de segmentación avanzados (YOLO-seg y SAM3) "
        "y se clasifica; la reconstrucción Image-to-3D (NeRF, Gaussian Splatting o difusión 3D) se deja "
        "explícitamente como trabajo futuro, una vez estabilizada la máscara 2D de calidad.",
    )
    p(
        doc,
        "La contribución práctica es el sistema CoArrCP IA: una aplicación web Flask que (i) corre inferencia "
        "YOLO-seg, (ii) propone puntos de conteo revisables, (iii) permite corrección y reentrenamiento, y "
        "(iv) integra SAM3 como asistente de máscara por texto o punto —evaluado en Google Colab por "
        "restricciones de peso del checkpoint e inestabilidad de red local.",
    )
    p(doc, "Objetivos específicos:", bold=True)
    bullets(
        doc,
        [
            "Demostrar un arreglo experimental reproducible de captura y organización de imágenes bentónicas.",
            "Implementar segmentación 2D avanzada (YOLO-seg entrenado + SAM3 foundation) alineada a Vía 1.",
            "Lograr detección/clasificación de componentes bentónicos con revisión humana tipo CoArrCP.",
            "Documentar el Manual del arreglo experimental (hardware, entorno, software, captura) adaptando la Guía rápida CoArrCP 2022r2 al nuevo stack.",
            "Justificar el uso de Colab para SAM3 (≈3.21 GB + red inestable en estación local).",
        ],
    )

    # ---------- 2 TRABAJOS RELACIONADOS ----------
    doc.add_heading("2. Trabajos relacionados", level=1)
    p(
        doc,
        "CoArrCP (Coberturas Arrecifales por Conteo de Puntos). Software clásico en Xojo/SQLite para "
        "estimar coberturas bentónicas por puntos sobre imágenes de transectos. Incluye catálogo de "
        "componentes (~112), editores de proyecto/transecto, verificación de BD y estadísticos "
        "(cobertura, GGMF/GMF, diversidad, complejidad topográfica). Esta tradición metodológica "
        "define el problema de usuario que CoArrCP IA automatiza parcialmente (Vargas Hernández, Guía rápida 2022r2).",
    )
    p(
        doc,
        "Segment Anything (SAM / SAM 2 / SAM 3). Familia de modelos foundation de Meta para segmentación "
        "promptable. SAM3 introduce conceptos abiertos por texto además de puntos/cajas, lo que lo hace "
        "atractivo para bentos donde el vocabulario (“coral”, “brain coral”, “rock”) es semántico y no "
        "solo geométrico (Kirillov et al.; Carion et al., SAM 3).",
    )
    p(
        doc,
        "YOLO de segmentación de instancias (Ultralytics YOLOv8/YOLO11-seg). Detectores en tiempo real "
        "que predicen máscaras y clases; adecuados para clasificación bentónica cuando hay dataset "
        "etiquetado (algas, corales, almejas, esponjas, arena).",
    )
    p(
        doc,
        "Image-to-3D (contexto del curso). NeRFs, 3D Gaussian Splatting, Zero123/TripoSR/Stable Fast 3D "
        "y MVS neuronal permiten reconstruir geometría desde máscaras/imágenes 2D. En este proyecto se "
        "reconocen como etapa posterior a una segmentación limpia (Vía 1 completa).",
    )
    p(
        doc,
        "Vía 2 (referencia). Segmentación nativa en point clouds (RANSAC+DBSCAN, PointNet++, MinkowskiEngine) "
        "con sensores RGB-D/ZED. No se implementó aquí; se documenta como alternativa para campañas con "
        "profundidad disponible.",
    )

    # ---------- 3 MATERIALES Y MÉTODOS ----------
    doc.add_heading("3. Materiales y métodos", level=1)

    doc.add_heading("3.1 Diseño metodológico (Vía 1)", level=2)
    p(
        doc,
        "Se adopta la Vía 1 del lineamiento del curso: segmentar en 2D con modelos avanzados y, "
        "posteriormente, habilitar reconstrucción 3D. El pipeline implementado es:",
    )
    bullets(
        doc,
        [
            "Captura → organización tipo CoArrCP (directorios de transecto / imágenes JPG).",
            "Etiquetado poligonal en Label Studio → COCO → conversión a YOLO-seg.",
            "Entrenamiento YOLO11n-seg (clases bentónicas) → inferencia en CoArrCP IA.",
            "Revisión humana de puntos (13 por imagen) + corrección de clases/nombres.",
            "Asistente SAM3 (texto/punto) en Colab para máscaras de alta calidad exportables a YOLO-seg.",
            "Reentrenamiento opcional desde correcciones (discos o máscaras SAM3).",
            "Trabajo futuro: Image-to-3D a partir de máscaras SAM3/YOLO (TripoSR / Gaussian Splatting).",
        ],
    )

    doc.add_heading("3.2 Datos", level=2)
    bullets(
        doc,
        [
            "Imágenes: fotografías submarinas JPG (ej. DSCN3095.JPG, 4608×3456).",
            "Conjunto etiquetado de desarrollo: 30 imágenes con polígonos en Label Studio (COCO).",
            "Clases YOLO-seg: algas, corales, almejas, esponjas, arena.",
            "Catálogo CoArrCP histórico: hasta ~112 componentes bentónicos (referencia metodológica; en IA se opera con subconjunto + clases libres).",
        ],
    )

    doc.add_heading("3.2.1 Obtención de las imágenes en campo (evidencia y protocolo)", level=3)
    p(
        doc,
        "Esta subsección documenta cómo se obtuvieron las fotografías del dataset a partir de "
        "metadatos EXIF embebidos en los archivos (no inventados) y del marco metodológico de "
        "monitoreo bentónico por fotografías / conteo de puntos alineado a CoArrCP y a protocolos "
        "internacionales de foto-cuadrantes y foto-transectos.",
    )

    p(doc, "A) Evidencia instrumental recuperada de los archivos (EXIF)", bold=True)
    p(
        doc,
        "Se inspeccionaron las 30 imágenes JPG del directorio de datos "
        "(ENTREGA_TESIS_SEGMENTACION_BENTONICA/datos/images/). Todas comparten la misma jornada "
        "de captura y dos cuerpos de cámara compacta impermeable Nikon, coherentes con trabajo "
        "de buceo sin housing externo:",
    )
    add_table(
        doc,
        ["Parámetro EXIF", "Valor observado en el dataset"],
        [
            ("Número de imágenes", "30 JPG"),
            ("Fecha DateTimeOriginal", "2021-06-08 (las 30)"),
            ("Ventana horaria", "09:45:14 → 11:59:23 (≈ 2 h 14 min de sesión)"),
            ("Cámaras", "Nikon COOLPIX AW130 (14 imgs) y Nikon COOLPIX W300 (16 imgs)"),
            ("Resolución", "4608 × 3456 px (RGB) en todas"),
            ("Óptica típica", "Focal 4.3 mm; f/2.8; ISO 125 (valores recurrentes)"),
            ("Flash", "15 con disparo de flash (código 9) y 15 sin flash (código 16)"),
            ("Software de cámara", "Firmware COOLPIX AW130 V1.0 / COOLPIX W300 V1.4"),
            ("GPS en EXIF", "No presente → el sitio geográfico no se afirma desde metadatos"),
        ],
        widths=[2.4, 4.3],
    )
    p(
        doc,
        "Interpretación: la secuencia temporal continua en la mañana del 8 de junio de 2021, "
        "con dos cámaras sumergibles operando en paralelo (o en relevo), es compatible con una "
        "inmersión / campaña fotográfica bentónica de duración limitada —no con un muestreo "
        "esporádico de laboratorio. El uso mixto de flash responde a la pérdida de color/luz "
        "con la profundidad, práctica documentada en protocolos de foto-cuadrantes que recomiendan "
        "flash interno en sitios más profundos o con menor iluminación (p. ej. implementaciones "
        "tipo GCRMN-Caribbean / fotoquadrat con cámara compacta).",
    )

    p(doc, "B) Requisitos de equipo según fabricante (lo que el buzo necesita llevar)", bold=True)
    p(
        doc,
        "Las cámaras identificadas en EXIF son modelos outdoor/underwater de Nikon. Según los "
        "manuales de referencia del fabricante:",
    )
    bullets(
        doc,
        [
            "Nikon COOLPIX AW130 y COOLPIX W300: impermeabilidad equivalente a JIS/IEC IPX8; disparo submarino hasta 30 m de profundidad durante un máximo de 60 minutos continuos (Nikon Reference Manual AW130 / W300).",
            "Rango de temperatura del agua declarado para uso submarino: 0–40 °C.",
            "No flotan: se requiere correa / sujeción al buzo para no perder el equipo.",
            "No abrir tapa de batería/tarjeta bajo el agua; enjuagar con agua dulce tras uso en mar.",
            "W300 incluye medidor de profundidad (rango aprox. 0–35 m) útil para bitácora de inmersión.",
            "Tarjeta de memoria, baterías cargadas y, en práctica de campo, cámara de respaldo (recomendación recurrente en manuales GCRMN de monitoreo fotográfico).",
        ],
    )
    p(
        doc,
        "Requisitos humanos y de seguridad (marco de monitoreo arrecifal con SCUBA/snorkel): "
        "mínimo dos observadores/buzos para seguridad mutua; certificación de buceo acorde a la "
        "profundidad de trabajo; boya/señalización de superficie según normativa local; pizarra "
        "submarina y cinta de transecto cuando el diseño sea foto-transecto (Hill & Wilkinson, 2004; "
        "manuales GCRMN).",
    )

    p(doc, "C) Procedimiento de toma fotográfica alineado a CoArrCP y foto-transecto", bold=True)
    p(
        doc,
        "CoArrCP fue diseñado precisamente para procesar fotografías (y material derivado de "
        "video-transectos) de coberturas bentónicas mediante conteo de puntos, con más de 15 años "
        "de aplicación en monitoreos del Parque Nacional Sistema Arrecifal Veracruzano "
        "(Vargas Hernández, Guía rápida CoArrCP 2022r2). El flujo de campo que alimenta ese "
        "análisis —y que explica la naturaleza de las imágenes DSCN* de este proyecto— es el "
        "siguiente (síntesis de la Guía CoArrCP + protocolos estándar de foto-cuadrante/foto-transecto):",
    )
    bullets(
        doc,
        [
            "Planificación en superficie: definir sitios/transectos; asignar códigos de carpeta (p. ej. TBA/TSO: las primeras tres letras identifican grupos de muestras en CoArrCP).",
            "Despliegue bajo el agua: los buzos colocan o siguen una línea de transecto (cinta métrica) preferentemente paralela a la cresta del arrecife y a profundidad aproximadamente constante (Hill & Wilkinson, 2004).",
            "Posición de la cámara: disparar perpendicular al sustrato, a distancia operativa controlada (típicamente del orden de ~0.5–1+ m según marco/escala), para reducir paralaje y mantener un área comparable entre fotos (GCRMN; SOP de photoquadrat).",
            "Muestreo a lo largo del transecto: fotografías sucesivas a intervalos regulares (o sobre cuadrantes fijos) que cubran el bentos sin sesgo deliberado hacia colonias “bonitas”.",
            "Iluminación: luz ambiente + flash cuando la profundidad o turbidez reduzcan color (coherente con el 50% de disparos con flash en este dataset).",
            "Registro: no borrar en campo; al salir, enjuague, descarga a disco, organización por transecto y verificación de integridad (CoArrCP exige nombres sin tildes/espacios problemáticos para SQLite).",
            "Análisis posterior (laboratorio): conteo de puntos (13 por defecto en CoArrCP) o, en este proyecto, etiquetado poligonal + YOLO-seg/SAM3 sobre el mismo tipo de fotografía de fondo.",
        ],
    )
    p(
        doc,
        "Lo que sí se afirma con evidencia: las 30 imágenes son capturas submarinas reales del "
        "08/06/2021 con cámaras impermeables Nikon AW130/W300 en sesión de ~2 h. "
        "Lo que no se afirma (ausencia de GPS/bitácora en el repositorio): nombre exacto del "
        "arrecife ni profundidad métrica por foto. El vínculo metodológico con CoArrCP/PNSAV se "
        "usa como marco de adquisición y análisis, no como georreferencia forzada de cada JPG.",
    )

    add_figure(
        doc,
        FIGS / "fig_original_DSCN3095.jpg",
        "Figura 1. Ejemplo del producto de campo: fotografía bentónica DSCN3095.JPG "
        "(Nikon, 4608×3456, 2021-06-08), típica del material que CoArrCP analiza por puntos "
        "y que este proyecto segmenta con YOLO-seg/SAM3.",
        width_in=5.0,
    )

    doc.add_heading("3.3 Software implementado", level=2)
    add_table(
        doc,
        ["Componente", "Rol", "Ubicación / notas"],
        [
            ("CoArrCP IA (Flask)", "UI de carga, inferencia, revisión de puntos, BD SQLite, reentreno", "coarcp_ia/; puerto 9001"),
            ("YOLO11n-seg", "Detección + segmentación + clasificación bentónica", "Ultralytics; pesos best.pt de corridas seg3"),
            ("SAM3 (Meta)", "Segmentación foundation por texto/punto", "Repo facebookresearch/sam3; checkpoint facebook/sam3 (gated)"),
            ("Notebook Colab", "Carga SAM3 en GPU T4, export máscaras/YOLO", "Copia_de_sam3.ipynb"),
            ("Scripts de entrenamiento", "prepare_coco_to_yolo_seg, train_seg, run_10_corridas", "Raíz / ENTREGA_TESIS/.../algoritmos"),
            ("API REST", "Puente futuro hacia cliente Xojo", "api_rest.py / documentación API_XOJO"),
        ],
    )

    # ---- MANUAL (OBLIGATORIO) ----
    doc.add_heading("3.4 Manual del arreglo experimental", level=2)
    p(
        doc,
        "Esta sección cumple el requisito estricto del curso. Adapta la Guía rápida CoArrCP 2022r2 "
        "(flujo Xojo + SQLite) al sistema CoArrCP IA + YOLO-seg + SAM3, preservando la lógica "
        "científica del conteo de puntos y la organización por transectos.",
    )

    doc.add_heading("3.4.1 Hardware", level=3)
    add_table(
        doc,
        ["Elemento", "Especificación / uso"],
        [
            ("Cámara de campo", "Nikon COOLPIX AW130 y COOLPIX W300 (EXIF del dataset); impermeables IPX8 hasta 30 m / 60 min (manuales Nikon). Resolución 4608×3456."),
            ("Estación de análisis local", "Apple Silicon (Mac) con Python 3.13, MPS para YOLO; 16+ GB RAM recomendados."),
            ("Estación SAM3 (nube)", "Google Colab con GPU NVIDIA T4 (o superior), CUDA 12.x/13.x, ~15 GB VRAM suficiente para inferencia imagen."),
            ("Almacenamiento", "SSD local para imágenes/transectos; cache HF (~3.21 GB sam3.pt) preferentemente en Colab o disco con red estable."),
            ("Red", "Local: Wi‑Fi doméstica inestable para descargas multi‑GB. Colab: backbone de Google (descarga checkpoint viable)."),
        ],
    )

    doc.add_heading("3.4.2 Entorno físico de captura", level=3)
    bullets(
        doc,
        [
            "Medio: fondo arrecifal submarino (fotografías bentónicas de sustrato; evidencia visual + EXIF de cámaras underwater).",
            "Jornada documentada por EXIF: 2021-06-08, 09:45–11:59 (sesión matutina ≈ 2 h).",
            "Iluminación: luz natural + flash en el 50% de los disparos (15/30), coherente con variación de profundidad/turbidez.",
            "Distancia y orientación: según protocolo de foto-transecto/foto-cuadrante — cámara aproximadamente perpendicular al sustrato (Hill & Wilkinson, 2004; GCRMN).",
            "Organización post-captura heredada de CoArrCP: carpetas por transecto; primeras tres letras como código (ej. TBA/TSO).",
            "Restricciones de nombres (Guía CoArrCP): evitar caracteres problemáticos para SQLite (tildes, espacios, guiones excesivos).",
        ],
    )

    doc.add_heading("3.4.3 Software y entorno de cómputo", level=3)
    add_table(
        doc,
        ["Capa", "Detalle"],
        [
            ("SO local", "macOS (Darwin) — desarrollo y CoArrCP IA"),
            ("SO Colab", "Ubuntu en VM Google (runtime GPU)"),
            ("Python", "3.12 (Colab) / 3.13 (venv local del proyecto)"),
            ("Frameworks", "PyTorch + CUDA (Colab); Ultralytics YOLO; OpenCV; Flask; NumPy"),
            ("SAM3", "pip install -e . desde github.com/facebookresearch/sam3; HF hub gated facebook/sam3"),
            ("Autenticación HF", "Token Read en Secrets de Colab (HF_TOKEN); usuario con acceso concedido (ArathW)"),
            ("BD aplicación", "SQLite en coarcp_ia/data/ (imágenes, puntos, inferencias, jobs de reentreno)"),
            ("Legacy CoArrCP", "Xojo + SQLite (Guía 2022r2) — referencia de UX/metodología, no runtime de este entregable"),
        ],
    )
    p(doc, "Dependencias mínimas CoArrCP IA (requirements.txt): flask≥3, ultralytics≥8, opencv-python≥4.8, numpy≥1.24, werkzeug≥3.", bold=False)

    doc.add_heading("3.4.4 Procedimiento de captura y análisis (reproducible)", level=3)
    p(doc, "A) Organización tipo CoArrCP (campo → disco)", bold=True)
    bullets(
        doc,
        [
            "Crear directorio de trabajo del proyecto (equivalente a la carpeta de BD CoArrCP).",
            "Crear subcarpetas por transecto/muestra; copiar JPG respetando orden alfabético de análisis.",
            "Opcional: registrar metadatos de proyecto (puntos por imagen=13, dimensiones en metros si se estimará densidad).",
        ],
    )
    p(doc, "B) Pipeline de etiquetado y entrenamiento YOLO-seg", bold=True)
    bullets(
        doc,
        [
            "Etiquetar polígonos en Label Studio → export COCO (result_coco.json).",
            "python3 prepare_coco_to_yolo_seg.py --coco-json result_coco.json --images-dir images --output-dir benthic_yolo_seg --class-order algas corales almejas esponjas arena",
            "python3 train_seg.py --data benthic_yolo_seg/dataset.yaml --model yolo11n-seg.pt --epochs 250 --imgsz 640 --batch 2 --device mps (o cuda)",
            "Seleccionar weights best.pt de la mejor corrida (p. ej. experimentos_seg3/runs/seg3_r012/weights/best.pt).",
        ],
    )
    p(doc, "C) Operación CoArrCP IA (equivalente moderno al panel «Análisis de imágenes»)", bold=True)
    bullets(
        doc,
        [
            "source .venv/bin/activate && python3 run.py --weights <best.pt> --host 127.0.0.1 --port 9001",
            "Abrir http://127.0.0.1:9001/ → cargar imágenes/carpetas (drag-and-drop).",
            "Ejecutar inferencia YOLO-seg; revisar 13 puntos (verde/rojo), arrastrar, renombrar, asignar componente.",
            "Marcar revisado; usar reentrenamiento desde correcciones si se desea mejorar el modelo.",
            "Atajos: J/K navegar, R revisado, A agregar punto, S máscara SAM3 (cuando el checkpoint esté disponible localmente).",
        ],
    )
    p(doc, "D) Asistente SAM3 en Google Colab (equivalente a aislar el objeto antes de 3D)", bold=True)
    bullets(
        doc,
        [
            "Runtime → GPU. Secret HF_TOKEN. Ejecutar Copia_de_sam3.ipynb.",
            "Clonar facebookresearch/sam3; descargar sam3.pt (~3.21 GB).",
            "Subir imagen coralina; prompt texto (ej. “coral”) y/o punto manual [x,y].",
            "Descargar sam3_outputs.zip (máscara PNG + overlay + .txt YOLO-seg) e importar a dataset/reentreno.",
        ],
    )

    doc.add_heading("3.4.5 Correspondencia Guía CoArrCP (Xojo) → CoArrCP IA", level=3)
    add_table(
        doc,
        ["Función en Guía rápida CoArrCP 2022r2", "Equivalente en CoArrCP IA / este proyecto"],
        [
            ("Editor de proyectos + Nueva BD SQLite", "Proyecto web + SQLite en coarcp_ia/data/; carga de lotes de imágenes"),
            ("Añadir transectos / carpetas de muestras", "Drag-and-drop de carpetas; listado de imágenes en UI"),
            ("Editor de componentes bentónicos (~112)", "Catálogo de componentes en BD + alta libre de clases; clases YOLO base (5)"),
            ("Análisis de imágenes: 13 puntos + lista de componentes", "Vista Revisión: 13 puntos IA + edición manual + atajos"),
            ("Actualizar registro / semáforo rojo-verde", "Estado revisado/pendiente por imagen"),
            ("Verificar BD / ubicar registro", "Consultas SQLite de puntos e inferencias; trazabilidad por image_id"),
            ("Estadística de coberturas / GGMF", "Cobertura por clase desde máscaras YOLO; estadísticos legacy siguen siendo referencia científica"),
            ("Exportar a Word/Excel (copiar-pegar)", "Export máscaras/labels; reportes técnicos DOCX generados por script"),
        ],
        widths=[2.8, 3.9],
    )

    doc.add_heading("3.5 Justificación del uso de Google Colab para SAM3", level=2)
    p(
        doc,
        "Las pruebas de SAM3 se ejecutaron en Google Colab (no en la estación macOS de desarrollo) por dos "
        "razones técnicas concurrentes:",
    )
    bullets(
        doc,
        [
            "Peso del checkpoint: sam3.pt ≈ 3.21 GB (gated en Hugging Face). En red local inestable la descarga se interrumpía (errores DNS/ConnectError, archivos incomplete a 0 B o progreso perdido tras cortes), impidiendo completar el arreglo experimental de inferencia.",
            "Aceleración: Colab ofrece GPU CUDA (T4) con anchos de banda de descarga típicos de 50–150 MB/s hacia el Hub, haciendo viable autenticar, bajar pesos y correr build_sam3_image_model + predict_inst / set_text_prompt en minutos.",
            "Reproducibilidad pedagógica: el notebook Copia_de_sam3.ipynb documenta el entorno cloud (CUDA, HF_TOKEN, HF_HUB_DISABLE_XET=1) como parte del Manual, separando claramente “estación YOLO local” vs “estación foundation SAM3 en nube”.",
        ],
    )
    p(
        doc,
        "Esta decisión no altera la Vía 1: la segmentación sigue siendo 2D avanzada; Colab es solo el "
        "arreglo de cómputo para el modelo foundation. Los artefactos (máscaras) se reintegran al flujo "
        "local CoArrCP IA / dataset YOLO.",
    )

    doc.add_heading("3.6 Detección y clasificación final", level=2)
    p(
        doc,
        "Conforme a los lineamientos (Detección 2D proyectada), YOLO-seg aporta bounding boxes implícitos "
        "y máscaras por clase. Los puntos CoArrCP se proyectan sobre la imagen; al moverse, pueden "
        "reconsultar la clase YOLO o aceptar etiqueta manual. SAM3 aporta máscaras de instancia de alta "
        "calidad (texto/punto) que refuerzan el aislamiento del sujeto antes de cualquier etapa 3D.",
    )

    # ---------- 4 RESULTADOS ----------
    doc.add_heading("4. Resultados", level=1)

    doc.add_heading("4.1 Entrenamiento YOLO-seg y sistema CoArrCP IA", level=2)
    bullets(
        doc,
        [
            "Pipeline COCO→YOLO-seg operativo; múltiples corridas documentadas (run_10_corridas_seg.py).",
            "Corrida de referencia ilustrada: experimentos_seg3/runs/seg3_r012 (pesos best.pt usados en CoArrCP IA).",
            "Aplicación Flask en localhost:9001 con revisión de puntos, SQLite y reentrenamiento desde correcciones.",
            "Integración de asistente SAM3 en código (sam3_assistant.py); ejecución práctica validada en Colab.",
        ],
    )
    p(
        doc,
        "Figura de contexto de campo: ver Figura 1 (§3.2.1). A continuación, métricas y predicciones de la corrida YOLO-seg de referencia.",
    )

    p(doc, "Curvas de entrenamiento y validación de la corrida seg3_r012:", bold=False)
    add_figure(
        doc,
        FIGS / "fig_yolo_results.png",
        "Figura 2. Resultados de entrenamiento YOLO11n-seg (seg3_r012): pérdidas box/seg/cls y métricas mAP.",
        width_in=6.2,
    )
    add_figure(
        doc,
        FIGS / "fig_yolo_mask_pr.png",
        "Figura 3. Curva Precision–Recall de máscaras (MaskPR) por clase bentónica — seg3_r012.",
        width_in=5.5,
    )
    add_figure(
        doc,
        FIGS / "fig_yolo_confusion.png",
        "Figura 4. Matriz de confusión normalizada (segmentación) — seg3_r012.",
        width_in=5.2,
    )

    p(doc, "Comparación etiquetas vs predicción en lote de validación:", bold=False)
    add_figure(
        doc,
        FIGS / "fig_yolo_val_labels.jpg",
        "Figura 5. Validación YOLO-seg — ground truth (val_batch0_labels).",
        width_in=5.5,
    )
    add_figure(
        doc,
        FIGS / "fig_yolo_val_pred.jpg",
        "Figura 6. Validación YOLO-seg — predicciones del modelo (val_batch0_pred).",
        width_in=5.5,
    )

    p(doc, "Inferencia sobre imágenes del dataset (predict):", bold=False)
    add_figure(
        doc,
        FIGS / "fig_yolo_pred_DSCN3095.jpg",
        "Figura 7. Inferencia YOLO-seg sobre DSCN3095 (máscaras + clases).",
        width_in=5.4,
    )
    add_figure(
        doc,
        FIGS / "fig_yolo_pred_DSCN3074.jpg",
        "Figura 8. Inferencia YOLO-seg sobre DSCN3074.",
        width_in=5.4,
    )
    add_figure(
        doc,
        FIGS / "fig_yolo_pred_DSCN3092.jpg",
        "Figura 9. Inferencia YOLO-seg sobre DSCN3092.",
        width_in=5.4,
    )

    doc.add_heading("4.2 Experimento SAM3 en Colab — imagen DSCN3095.JPG", level=2)
    p(
        doc,
        "Configuración: GPU NVIDIA (CUDA) en Google Colab, usuario HF ArathW con acceso gated, "
        "checkpoint sam3.pt 3.21 GB OK, modelo con enable_inst_interactivity=True.",
    )
    add_table(
        doc,
        ["Modo", "Configuración", "Hallazgo cuantitativo", "Interpretación"],
        [
            (
                "Texto",
                'TEXT_PROMPT = "coral"',
                "23 instancias detectadas; máscara exportada top-1: 11.5% del área (1 826 492 px), 1 componente, 1 polígono YOLO (32 vértices). BBox ≈ (1171,1807)–(2772,3286).",
                "Segmentación limpia de una colonia coralina; apta como seed de etiqueta/reentreno.",
            ),
            (
                "Punto",
                "Fallback centro (2304, 1728); score 0.621",
                "Cobertura 38.8% (6 181 327 px), 48 componentes, 62 polígonos YOLO; IoU vs texto = 0.00.",
                "El centro cayó en sustrato/roca, no en el coral. Confirma que el punto debe ser guiado (clic/coords sobre el organismo).",
            ),
        ],
        widths=[1.1, 1.6, 2.2, 2.0],
    )

    add_figure(
        doc,
        FIGS / "fig_sam3_text_overlay.jpg",
        "Figura 10. SAM3 — overlay del mejor resultado por texto (“coral”) sobre DSCN3095.",
        width_in=5.2,
    )
    add_figure(
        doc,
        FIGS / "fig_sam3_text_preview.jpg",
        "Figura 11. SAM3 texto — vista dual (overlay | máscara binaria top-1). Cobertura ≈ 11.5%.",
        width_in=5.8,
    )
    add_figure(
        doc,
        FIGS / "fig_sam3_point_overlay.jpg",
        "Figura 12. SAM3 — overlay del modo punto en el centro de la imagen (score 0.621).",
        width_in=5.2,
    )
    add_figure(
        doc,
        FIGS / "fig_sam3_point_preview.jpg",
        "Figura 13. SAM3 punto — vista dual. La máscara cubre sustrato central (~38.8%), no el coral.",
        width_in=5.8,
    )
    add_figure(
        doc,
        FIGS / "fig_sam3_compare.jpg",
        "Figura 14. Comparación espacial texto vs punto (IoU = 0): regiones disjuntas.",
        width_in=5.0,
    )

    p(
        doc,
        "Conclusión experimental: el prompt semántico de SAM3 es efectivo para aislar coral en escenas "
        "complejas; el modo punto es poderoso pero sensible a la ubicación del prompt. Para CoArrCP, "
        "la combinación recomendada es: YOLO para clasificación masiva + SAM3 texto/punto dirigido para "
        "correcciones de máscara de alta calidad. La Figura 7 (YOLO) y las Figuras 10–11 (SAM3 texto) "
        "sobre la misma DSCN3095 ilustran ambas vías de segmentación 2D de la Vía 1.",
    )

    doc.add_heading("4.3 Cumplimiento de la rúbrica (autoevaluación documental)", level=2)
    add_table(
        doc,
        ["Criterio (pts)", "Evidencia en el proyecto", "Nivel orientativo"],
        [
            ("Arreglo experimental (10)", "Captura JPG, organización transectos, Manual 3.4, Colab GPU documentado", "Bueno–Excelente"),
            ("Procesamiento/Segmentación Vía 1 (20)", "YOLO-seg + SAM3; máscara texto limpia; punto centro documentado como fallo controlado", "Bueno (ruido residual en punto no guiado)"),
            ("Detección/Clasificación (20)", "5 clases YOLO + puntos CoArrCP revisables; tiempos locales aceptables en MPS/CUDA", "Bueno"),
            ("Estructura reporte (15)", "Secciones 1–6 obligatorias presentes", "Excelente (si se completa citación formal)"),
            ("Manual arreglo (15)", "§3.4 Hardware/Entorno/Software/Procedimiento + puente Guía Xojo", "Excelente"),
        ],
    )
    p(
        doc,
        "Nota: la presentación oral (20 pts) queda fuera del alcance de este entregable por decisión del autor.",
    )

    # ---------- 5 CONCLUSIONES ----------
    doc.add_heading("5. Conclusiones", level=1)
    bullets(
        doc,
        [
            "Se demostró un ciclo 2D completo alineado a la Vía 1 del curso: captura → segmentación avanzada → clasificación → revisión humana tipo CoArrCP.",
            "CoArrCP IA moderniza el flujo de la Guía rápida 2022r2 sin romper la lógica científica del conteo de puntos y el catálogo bentónico.",
            "SAM3 en Colab resolvió el cuello de botella de peso/red: el checkpoint de 3.21 GB y la red local inestable hacían inviable la prueba en Mac; en T4 la inferencia y exportación fueron exitosas.",
            "En DSCN3095, el prompt “coral” isoló correctamente una colonia (11.5% área, máscara limpia); el punto automático al centro no es un proxy válido y debe evitarse en producción.",
            "La reconstrucción 3D (NeRF/GS/difusión) no se ejecutó en esta fase; queda legitimada como siguiente eslabón de la Vía 1 una vez fijadas máscaras confiables.",
        ],
    )

    # ---------- 6 TRABAJO FUTURO ----------
    doc.add_heading("6. Trabajo futuro", level=1)
    bullets(
        doc,
        [
            "Image-to-3D: alimentar TripoSR / Stable Fast 3D / Gaussian Splatting con recortes enmascarados por SAM3.",
            "Exportar N mejores instancias de texto SAM3 (no solo top-1) hacia el dataset de reentrenamiento.",
            "Cache local de sam3.pt en disco estable o mirror institucional para dejar de depender de Colab.",
            "Evaluación formal mAP/IoU YOLO-seg vs máscaras SAM3 en un split hold-out.",
            "Puente Xojo↔API REST para que biólogos conserven el cliente histórico con backend IA.",
            "Explorar Vía 2 (ZED 2i + RANSAC/DBSCAN o PointNet++) en campañas con profundidad.",
        ],
    )

    doc.add_heading("Referencias", level=1)
    bullets(
        doc,
        [
            "Vargas Hernández, J. M. Guía rápida CoArrCP 2022r2. Coberturas Arrecifales por Conteo de Puntos.",
            "Nikon Corporation. COOLPIX AW130 Reference Manual (impermeabilidad IPX8, 30 m / 60 min).",
            "Nikon Corporation. COOLPIX W300 Reference Manual (IPX8, 30 m / 60 min; medidor de profundidad).",
            "Hill, J. & Wilkinson, C. (2004). Methods for Ecological Monitoring of Coral Reefs. Australian Institute of Marine Science (AIMS), Townsville.",
            "GCRMN / manuals regionales de monitoreo arrecifal: uso de foto-transectos y foto-cuadrantes; cámara perpendicular al sustrato; registro permanente de imágenes.",
            "Kirillov, A. et al. Segment Anything. ICCV / Meta AI.",
            "Carion, N. et al. SAM 3: Segment Anything with Concepts. Meta Superintelligence Labs.",
            "Jocher, G. et al. Ultralytics YOLO (segmentation). https://github.com/ultralytics/ultralytics",
            "Mildenhall, B. et al. NeRF: Representing Scenes as Neural Radiance Fields. ECCV 2020.",
            "Kerbl, B. et al. 3D Gaussian Splatting for Real-Time Radiance Field Rendering. SIGGRAPH 2023.",
            "Facebook Research. sam3 (código y checkpoints gated). https://github.com/facebookresearch/sam3",
            "Metadatos EXIF del dataset propio (30× DSCN*.JPG, DateTimeOriginal=2021-06-08) — fuente primaria de la sección 3.2.1.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("DOCX:", OUT)
    return OUT


def build_md():
    """Copia resumida en Markdown para edición rápida."""
    text = f"""# Reporte Técnico — Proyecto Final
## Procesamiento y reconocimiento visual en 3D
### CoArrCP IA + YOLO-seg + SAM3 (Vía 1)

**Autor:** Wilder Arath Utrera Vargas  
**Fecha:** Julio 2026  
**Documento Word:** `{OUT.name}`

> Este Markdown es espejo breve. El entregable completo con tablas y Manual §3.4 está en el DOCX.

## Cumplimiento de lineamientos
- **Vía 1 (2D→3D):** segmentación SAM3 + YOLO-seg; Image-to-3D = trabajo futuro.
- **Detección/clasificación:** YOLO-seg + puntos CoArrCP revisables.
- **Manual del arreglo experimental:** incluido en DOCX §3.4 (adaptado de Guía rápida CoArrCP 2022r2).
- **Sin presentación oral** en esta entrega.

## Por qué Colab para SAM3
1. Checkpoint `sam3.pt` ≈ **3.21 GB** (gated HF).
2. Red local **inestable** (cortes DNS/ConnectError; descargas incomplete).
3. Colab GPU T4: descarga rápida + inferencia CUDA reproducible (`Copia_de_sam3.ipynb`).

## Resultado clave DSCN3095
| Modo | Resultado |
|------|----------|
| Texto `"coral"` | 23 instancias; top-1 limpia **11.5%** área |
| Punto centro | score 0.621 pero **38.8%** sustrato; **IoU=0** vs texto |

## Estructura DOCX
1. Introducción  
2. Trabajos relacionados  
3. Materiales y métodos (+ **Manual 3.4**)  
4. Resultados  
5. Conclusiones  
6. Trabajo futuro  
"""
    MD_OUT.write_text(text, encoding="utf-8")
    print("MD:", MD_OUT)


if __name__ == "__main__":
    build()
    build_md()
