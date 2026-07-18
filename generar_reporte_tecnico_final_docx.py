#!/usr/bin/env python3
"""
Reporte Técnico Final DCU — integra documentos previos del proyecto de tesis:
- Evaluacion_Heuristica_CoArrCP_IA.docx
- Informe_Biologo_Interfaz_y_Atajos_CoArrCP_IA.docx
- Validacion_Experto_Prototipo.docx
- Metodologia del protecto_de_tesis_sobre_SegmentacionDeImagenesBentonicas.pdf
- coarcp_ia/ATAJOS_TECLADO.md, README_pipeline, API_XOJO_COARRCP_IA.md
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


DOCS = Path(__file__).resolve().parent / "ENTREGA_TESIS_SEGMENTACION_BENTONICA" / "documentos"
OUTPUT = DOCS / "Reporte_Tecnico_Final_DCU_CoArrCP_IA.docx"
PRESENTATION = DOCS / "Guion_Presentacion_Final_DCU.docx"


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# --- Datos tomados de Evaluacion_Heuristica_CoArrCP_IA.docx ---
EVAL1 = [
    ("1", "Visibilidad del estado del sistema", "2", "Resultados de IA visibles; sin estado claro del modelo ni progreso."),
    ("2", "Correspondencia con el mundo real", "2", "Clases bentónicas correctas; conf/iou sin explicación en pantalla."),
    ("3", "Control y libertad del usuario", "1", "Sin mover puntos, marcar revisado ni corregir inferencias."),
    ("4", "Consistencia y estándares", "2", "Interfaz distinta al flujo CoArrCP esperado."),
    ("5", "Prevención de errores", "1", "Arranque por terminal; servidor no siempre activo."),
    ("6", "Reconocimiento antes que recuerdo", "2", "Sin menú ni recordatorio del flujo."),
    ("7", "Flexibilidad y eficiencia", "1", "Solo carga archivo a archivo."),
    ("8", "Diseño estético y minimalista", "2", "Página simple sin jerarquía visual clara."),
    ("9", "Diagnóstico y recuperación de errores", "1", "Errores poco claros sin guía."),
    ("10", "Ayuda y documentación", "1", "Sin ayuda sobre parámetros ni puntos de conteo."),
]

EVAL2 = [
    ("1", "Visibilidad del estado del sistema", "3", "Pill «Modelo cargado», métricas, flash, estado al guardar puntos."),
    ("2", "Correspondencia con el mundo real", "2", "Flujo CoArrCP; conf/iou aún técnicos."),
    ("3", "Control y libertad del usuario", "3", "Puntos arrastrables, revisado/pendiente, carpetas, drag-and-drop."),
    ("4", "Consistencia y estándares", "3", "Menú lateral con etapas del flujo bentónico."),
    ("5", "Prevención de errores", "2", "Validación al subir; aún requiere terminal (run.py)."),
    ("6", "Reconocimiento antes que recuerdo", "3", "Menú visible, leyenda verde/rojo, tabla de puntos."),
    ("7", "Flexibilidad y eficiencia", "3", "Lotes, drag-and-drop, atajos J/K/R/D."),
    ("8", "Diseño estético y minimalista", "3", "Modo revisión con foto dominante; parámetros plegados."),
    ("9", "Diagnóstico y recuperación de errores", "2", "Flash útiles; faltan guías si falla YOLO."),
    ("10", "Ayuda y documentación", "2", "README, ATAJOS_TECLADO.md, diálogo ?; sin tooltips en UI."),
]

CAMBIOS_HEURISTICA = [
    "Menú lateral CoArrCP (inicio, carga, inferencia, revisión, componentes, BD, estadística, config).",
    "Drag-and-drop de imágenes y carpetas (coarcp_ia/templates/index.html).",
    "13 puntos verdes/rojos editables; recálculo de clase al soltar (API /imagen/<id>/puntos).",
    "SQLite: imágenes, inferencias, cobertura, detecciones, puntos (coarcp_ia/data/).",
    "Parámetros avanzados plegados (conf, iou, puntos) — heurística minimalista.",
    "Atajos Shneiderman: J/K/R/U/D/G/L/F/?/Esc (ATAJOS_TECLADO.md).",
    "Modo revisión: fotografía dominante; máscaras limpias sin artefactos morados de YOLO.",
    "run.py como único punto de arranque; API REST para futuro Xojo (API_XOJO_COARRCP_IA.md).",
    "Corrección escala máscaras (480×640 → resolución original) tras error broadcast NumPy.",
    "Botón Menú + Esc para volver al inicio; paneles G/L corregidos en modo revisión.",
]

# Valoraciones del Informe_Biologo_Interfaz_y_Atajos_CoArrCP_IA.docx
VALORACION_BIOLOGO = [
    ("Velocidad para revisar una foto", "4/5", "K y R agilizan; arrastrar puntos sigue manual"),
    ("Claridad de la interfaz", "4/5", "Modo foco útil; barra superior podría ser más discreta"),
    ("Utilidad de teclas rápidas", "5/5", "Esenciales en lotes grandes"),
    ("Confianza en segmentación IA", "3–4/5", "Siempre revisa puntos; IA no sustituye criterio"),
    ("¿Uso en campaña real?", "Sí", "Como apoyo, no sustituto del biólogo"),
]


def build_report():
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # PORTADA
    doc.add_heading("Reporte Técnico Final", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Diseño Centrado en el Usuario — Sistema de tesis").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for label, val in [
        ("Nombre completo:", "[COMPLETAR NOMBRE]"),
        ("Matrícula:", "[COMPLETAR MATRÍCULA]"),
        ("Sistema:", "CoArrCP IA — Segmentación automatizada de componentes bentónicos"),
        ("Tesis:", "Segmentación de imágenes bentónicas con YOLO-seg e integración CoArrCP"),
        ("Fecha:", "Junio 2026"),
    ]:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(val)

    doc.add_heading("Corpus documental de referencia (documentos anteriores)", level=2)
    add_table(
        doc,
        ["Documento", "Contenido integrado en este reporte"],
        [
            ("Evaluacion_Heuristica_CoArrCP_IA.docx", "Secciones 5 y 7; tablas Eval. 1 y 2; comparativa 1.6→2.7"),
            ("Informe_Biologo_Interfaz_y_Atajos_CoArrCP_IA.docx", "Secciones 2, 3, 6; perfil, flujo, valoraciones"),
            ("Validacion_Experto_Prototipo.docx", "Sección 6; protocolo de prueba piloto 15–20 min"),
            ("Metodologia...Bentonicas.pdf", "Sección 1; contexto Label Studio, 30 imgs, YOLO11n-seg"),
            ("ATAJOS_TECLADO.md", "Secciones 4 y 5; regla Shneiderman documentada"),
            ("API_XOJO_COARRCP_IA.md", "Sección 7; arquitectura objetivo Flask + Xojo"),
            ("README_pipeline.md / ENTREGA_TESIS README", "Sección 4; pipeline técnico reproducible"),
        ],
        [2.5, 3.8],
    )
    doc.add_page_break()

    # 1 INTRODUCCIÓN
    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(
        "El sistema de tesis automatiza la segmentación de componentes bentónicos en imágenes "
        "coralinas: algas, corales, almejas, esponjas y arena. El pipeline técnico documentado en "
        "Metodologia del protecto_de_tesis_sobre_SegmentacionDeImagenesBentonicas.pdf comprende: "
        "etiquetado en Label Studio (Brush Labels → COCO), conversión a YOLO-seg, entrenamiento "
        "con YOLO11n-seg (~30 imágenes, corrida seg3_r012, Mask mAP50 ≈ 0.53) e inferencia local."
    )
    doc.add_paragraph(
        "CoArrCP IA es el prototipo de interfaz para operar ese modelo en un flujo inspirado en "
        "CoArrCP: cargar fotos de campaña, ejecutar IA, revisar máscaras y 13 puntos de conteo, "
        "marcar imágenes como revisadas y exportar resultados. Está implementado en Python Flask "
        "(carpeta coarcp_ia/, arranque con run.py, puerto 9001). La integración con Xojo + SQLite "
        "de producción queda documentada como trabajo futuro (API_XOJO_COARRCP_IA.md)."
    )
    doc.add_paragraph(
        "Este reporte no describe solo el código actual: documenta el proceso DCU, las decisiones "
        "entre prototipos (app_flask simple → coarcp_ia_flask → CoArrCP IA integrado) y cómo la "
        "evaluación heurística y el informe del biólogo experto guiaron las iteraciones."
    )

    # 2 PERFIL
    doc.add_heading("2. Perfil de usuario", level=1)
    doc.add_paragraph(
        "Fuente: Plantilla 1 + Informe_Biologo_Interfaz_y_Atajos_CoArrCP_IA.docx (perspectiva de "
        "usuario experto en conteo bentónico y revisión de cuadrantes coralinos)."
    )
    doc.add_heading("Usuario primario — Biólogo marino experto", level=2)
    add_table(
        doc,
        ["Atributo", "Descripción (del informe del biólogo)"],
        [
            ("Rol", "Revisa cientos de fotografías bentónicas por campaña"),
            ("Contexto", "PC de escritorio; sesiones largas tras fototransectos en arrecife"),
            ("Motivación", "La imagen del fondo marino debe ser lo principal; sin menús técnicos"),
            ("Frustración inicial", "Terminal, parámetros conf/iou, fatiga por repetir clics"),
            ("Flujo ideal", "Arrastrar carpeta → procesar → revisar puntos → R → K → siguiente"),
            ("Criterio", "La IA apoya pero no sustituye su juicio taxonómico"),
        ],
        [1.8, 4.7],
    )
    doc.add_heading("Usuario secundario — Técnico / estudiante de tesis", level=2)
    doc.add_paragraph(
        "Carga lotes, entrena o cambia pesos del modelo (best.pt), mapea clases YOLO en pestaña "
        "Componentes (GGMF, GMF). Usa README_pipeline y run.py. Documentado en ENTREGA_TESIS README."
    )

    # 3 ESCENARIOS
    doc.add_heading("3. Escenarios y tareas", level=1)
    doc.add_paragraph(
        "Escenario tomado del Informe del biólogo y Validacion_Experto_Prototipo.docx."
    )
    doc.add_heading("Escenario A — Revisión post-campaña (principal)", level=2)
    doc.add_paragraph(
        "El biólogo tiene una carpeta con fotos JPG del día de muestreo. Debe verificar la "
        "segmentación automática, ajustar puntos sobre corales/algas y dejar el lote marcado "
        "como revisado para el informe CoArrCP."
    )
    add_table(
        doc,
        ["Tarea", "Acción", "Evidencia en prototipo", "Decisión DCU"],
        [
            ("T1", "Cargar carpeta", "Drag-and-drop en Cargar imágenes", "Reducir pasos; Nielsen eficiencia"),
            ("T2", "Procesar lote", "Botón «Procesar con IA»", "Una acción clara; ocultar conf/iou"),
            ("T3", "Verificar lote", "Pestaña Inferencia IA + SQLite", "Visibilidad del estado"),
            ("T4", "Revisar foto", "Revisión manual, foto grande", "Minimalismo: foto = jerarquía"),
            ("T5", "Mover puntos", "Overlay verde/rojo arrastrable", "Control y libertad (Eval. 1→3)"),
            ("T6", "Confirmar", "Tecla R o botón Revisada", "Shneiderman: atajos expertos"),
            ("T7", "Siguiente foto", "Tecla K / →", "Secuencia R→K documentada en informe biólogo"),
            ("T8", "Exportar", "Tecla D", "Descarga JPG con puntos dibujados"),
        ],
        [0.5, 1.5, 1.8, 2.7],
    )
    doc.add_heading("Escenario B — Validación con experto (15–20 min)", level=2)
    doc.add_paragraph(
        "Protocolo de Validacion_Experto_Prototipo.docx: cargar 3–5 imágenes, ajustar 2 puntos, "
        "usar R y K sin ratón, probar ? y F. Cuestionario sobre fotografía dominante, atajos y "
        "agilidad en tarea repetitiva (escala 1–5)."
    )

    # 4 PROCESO DE DISEÑO
    doc.add_heading("4. Proceso de diseño", level=1)
    etapas = [
        (
            "Etapa 0 — Investigación y metodología de tesis",
            "PDF de metodología: 30 imágenes coralinas, Label Studio, COCO, YOLO-seg, métricas "
            "de corridas (seg3_r012, etc.). El problema de usuario: etiquetado/revisión manual lenta.",
            "Definir que la UI debe apoyar revisión, no solo inferencia técnica.",
            "Metodologia...Bentonicas.pdf",
        ),
        (
            "Etapa 1 — Prototipo mínimo (app_flask.py original)",
            "Una página: subir imagen, ver segmentación y cobertura. Sin menú, sin SQLite, sin puntos.",
            "Validar que el modelo funciona; Evaluación heurística 1 (promedio 1.6/3).",
            "Evaluacion_Heuristica — Eval. 1",
        ),
        (
            "Etapa 2 — coarcp_ia_flask (media fidelidad)",
            "Carpeta descargada con menú lateral, SQLite, 13 puntos fijos, pestañas CoArrCP.",
            "Alinear flujo con software de referencia; aún parámetros visibles y sin atajos.",
            "README coarcp_ia_flask",
        ),
        (
            "Etapa 3 — Integración en proyecto de tesis (coarcp_ia/)",
            "Fusión con API Xojo, run.py, drag-and-drop, editor de puntos, informe biólogo.",
            "Evaluación 2 (promedio 2.7/3); prototipo de alta fidelidad operativo.",
            "Informe_Biologo + Evaluacion_Heuristica Eval. 2",
        ),
        (
            "Etapa 4 — Iteraciones DCU del curso",
            "Modo revisión minimalista, atajos Shneiderman, corrección puntos morados/duplicados "
            "(plot limpio vs overlay), fix broadcast máscaras, Menú/Esc, paneles G/L.",
            "Respuesta directa a hallazgos heurísticos y prueba piloto del biólogo.",
            "ATAJOS_TECLADO.md; commits de coarcp_ia/",
        ),
    ]
    for titulo, que, por_que, evidencia in etapas:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(que)
        p = doc.add_paragraph()
        p.add_run("Decisión: ").bold = True
        p.add_run(por_que)
        p2 = doc.add_paragraph()
        p2.add_run("Evidencia documental: ").bold = True
        p2.add_run(evidencia)
    doc.add_paragraph(
        "[Anexos B–D: insertar capturas de app_flask inicial, coarcp_ia_flask y CoArrCP IA actual "
        "en http://127.0.0.1:9001 — pestañas Cargar, Inferencia, Revisión con puntos.]"
    )

    # 5 EVALUACIÓN HEURÍSTICA
    doc.add_heading("5. Evaluación heurística", level=1)
    doc.add_paragraph(
        "Reproduce y resume Evaluacion_Heuristica_CoArrCP_IA.docx (metodología 5 pasos del curso, "
        "escala Likert 1–3). Evaluación 1 = app_flask simple. Evaluación 2 = CoArrCP IA actual."
    )
    doc.add_heading("5.1 Evaluación 1 — Prototipo original (promedio 1.6)", level=2)
    add_table(doc, ["#", "Heurística", "Nota", "Justificación"], EVAL1, [0.35, 2.0, 0.45, 3.5])
    doc.add_heading("5.2 Cambios realizados (Paso 3 del curso)", level=2)
    for c in CAMBIOS_HEURISTICA:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_heading("5.3 Evaluación 2 — Prototipo mejorado (promedio 2.7)", level=2)
    add_table(doc, ["#", "Heurística", "Nota", "Justificación"], EVAL2, [0.35, 2.0, 0.45, 3.5])
    doc.add_heading("5.4 Comparativa y hallazgos con severidad", level=2)
    add_table(
        doc,
        ["Heurística", "Eval.1", "Eval.2", "Severidad inicial", "Mejora / estado"],
        [
            ("Control y libertad", "1", "3", "Alta", "Corregido: puntos + R/U"),
            ("Flexibilidad y eficiencia", "1", "3", "Alta", "Corregido: atajos + lotes"),
            ("Prevención de errores", "1", "2", "Alta", "Parcial: run.py; falta empaquetado"),
            ("Ayuda y documentación", "1", "2", "Media", "Parcial: ? y .md; faltan tooltips"),
            ("Diseño minimalista", "2", "3", "Media", "Corregido: foto dominante, params ocultos"),
            ("Visibilidad del estado", "2", "3", "Media", "Corregido: pill modelo, métricas"),
            ("PROMEDIO", "1.6", "2.7", "—", "+1.1 puntos"),
        ],
        [1.8, 0.5, 0.5, 0.9, 2.8],
    )
    doc.add_paragraph(
        "Pendientes (trabajo futuro, también en Evaluacion_Heuristica): tooltips conf/iou, "
        "ayuda en Configuración, empaquetado sin terminal, integración Xojo."
    )

    # 6 PLAN DE PRUEBAS
    doc.add_heading("6. Plan de pruebas de usabilidad", level=1)
    doc.add_paragraph("Plantilla 2 + protocolo Validacion_Experto_Prototipo.docx.")
    doc.add_heading("6.1 Participantes", level=2)
    doc.add_paragraph(
        "Mínimo 3: (1) biólogo marino experto — perfil del Informe_Biologo; "
        "(2) técnico de laboratorio; (3) estudiante ecología marina."
    )
    doc.add_heading("6.2 Tareas del protocolo piloto (ya diseñadas)", level=2)
    for i, t in enumerate([
        "Cargar 3–5 imágenes arrastrando carpeta",
        "Abrir Revisión manual en primera imagen",
        "Ajustar ≥2 puntos sobre corales/algas",
        "Marcar revisada con R; siguiente con K sin ratón",
        "Descargar con D; abrir ayuda con ?; probar F",
    ], 1):
        doc.add_paragraph(f"{i}. {t}", style="List Number")
    doc.add_heading("6.3 Métricas", level=2)
    add_table(
        doc,
        ["Métrica", "Fuente", "Meta"],
        [
            ("Tiempo por imagen revisada", "Cronómetro", "< 90 s tras práctica"),
            ("Éxito mover punto y ver clase", "Observación", "100%"),
            ("Utilidad atajos (1–5)", "Cuestionario Validacion_Experto", "≥ 4"),
            ("Velocidad revisar foto (1–5)", "Informe biólogo (piloto simulado)", "4/5 obtenido"),
            ("SUS / satisfacción global", "Plantilla 2", "≥ 70 SUS"),
        ],
        [2.0, 2.2, 2.3],
    )
    doc.add_heading("6.4 Resultados del informe del biólogo (piloto documentado)", level=2)
    add_table(doc, ["Aspecto", "Resultado", "Comentario"], VALORACION_BIOLOGO, [2.0, 1.0, 3.5])
    doc.add_paragraph("Lo que funciona (Informe_Biologo):")
    for item in [
        "Foto del arrecife domina la pantalla",
        "conf/iou ocultos en Parámetros avanzados",
        "K y R para ritmo continuo",
        "Puntos verdes/rojos intuitivos",
        "Arrastrar puntos es natural para conteo bentónico",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("Mejoras solicitadas (Informe_Biologo + Validacion_Experto):")
    for item in [
        "Atajos para puntos 10–13",
        "Indicador «guardando» más visible",
        "Exportar lote completo (hoy solo D por imagen)",
        "Ayuda breve sobre conf/iou para colegas",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "[Espacio para completar tras prueba formal con 3 participantes reales y firmas en "
        "Validacion_Experto_Prototipo.docx.]"
    )

    # 7 REFLEXIÓN
    doc.add_heading("7. Reflexión final", level=1)
    doc.add_heading("¿Qué cambió en tu diseño a partir del curso?", level=2)
    doc.add_paragraph(
        "La evaluación heurística cuantificó la brecha entre «demo para el desarrollador» y "
        "«herramienta para el biólogo» (1.6 → 2.7). El informe del biólogo confirmó que la "
        "secuencia R→K y la jerarquía de la fotografía no son opcionales. Apliqué Nielsen "
        "(minimalismo, visibilidad), Shneiderman (atajos) e ISO 9241-210 (iterar con evidencia). "
        "Los documentos anteriores no son anexos decorativos: son la trazabilidad del diseño."
    )
    doc.add_heading("¿Qué harías diferente con más tiempo?", level=2)
    doc.add_paragraph(
        "Probar con 3 biólogos reales antes de codificar el editor de puntos; prototipo Figma "
        "del modo revisión; no depender de terminal en absoluto; exportación masiva; validar "
        "Plantilla 2 completa antes de la entrega de tesis."
    )
    doc.add_heading("¿Cómo continuará en la investigación de tesis?", level=2)
    doc.add_paragraph(
        "Integrar Flask con Xojo vía API (API_XOJO_COARRCP_IA.md); ampliar dataset (>30 imgs); "
        "métricas de tiempo manual vs IA+revisión; capítulo de tesis enlazando metodología ML "
        "y DCU; prueba SUS formal. Arquitectura objetivo: Xojo = UI + SQLite producción; "
        "Flask = motor IA (documentado en metodología PDF)."
    )

    # 8 REFERENCIAS
    doc.add_heading("8. Referencias", level=1)
    for r in [
        "ISO 9241-210:2019 — Diseño centrado en el usuario.",
        "Nielsen, J. (1994) — Heurísticas de usabilidad (documento Evaluacion_Heuristica).",
        "Shneiderman, B. — Atajos para usuarios expertos (ATAJOS_TECLADO.md).",
        "Norman, D. — Affordances y carga cognitiva (ocultar parámetros avanzados).",
        "Metodologia del protecto_de_tesis_sobre_SegmentacionDeImagenesBentonicas.pdf.",
        "Ultralytics YOLO11-seg — Documentación técnica del modelo.",
        "Material del curso DCU — Plantillas 1, 2, 3 y 4.",
    ]:
        doc.add_paragraph(r, style="List Bullet")

    # ANEXOS
    doc.add_heading("Anexos — documentos anteriores del proyecto", level=1)
    for a in [
        "Anexo A — Plantilla 1: Perfil de usuario (sección 2 + Informe_Biologo)",
        "Anexo B — Evaluacion_Heuristica_CoArrCP_IA.docx (evaluación completa)",
        "Anexo C — Informe_Biologo_Interfaz_y_Atajos_CoArrCP_IA.docx",
        "Anexo D — Validacion_Experto_Prototipo.docx (protocolo y cuestionario)",
        "Anexo E — Plantilla 2: Plan de pruebas (sección 6)",
        "Anexo F — ATAJOS_TECLADO.md",
        "Anexo G — API_XOJO_COARRCP_IA.md",
        "Anexo H — Metodologia del protecto_de_tesis_sobre_SegmentacionDeImagenesBentonicas.pdf",
        "Anexo I — Capturas CoArrCP IA (Cargar / Inferencia / Revisión) [ADJUNTAR]",
        "Anexo J — Bocetos en papel o wireframes [ADJUNTAR SI APLICA]",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    doc.save(OUTPUT)
    print(f"Reporte: {OUTPUT}")


def build_presentation():
    doc = Document()
    doc.add_heading("Guion presentación DCU (15 min + 5 min Preguntas)", 0)
    doc.add_paragraph(
        "Basado en Reporte_Tecnico_Final_DCU y documentos: Evaluacion_Heuristica, "
        "Informe_Biologo, Validacion_Experto."
    )
    blocks = [
        ("1. Portada (30 s)", "Nombre, matrícula, CoArrCP IA, tesis segmentación bentónica YOLO-seg."),
        ("2. Usuario y problema (2 min)", "Biólogo, cientos de fotos, fatiga. Citar valoración 5/5 en atajos (Informe_Biologo)."),
        ("3. Proceso (4 min)", "app_flask (1.6) → coarcp_ia_flask → CoArrCP IA (2.7). Mostrar tabla comparativa."),
        ("4. Prototipo actual (4 min)", "Demo: drag-and-drop, revisión, punto verde/rojo, R, K. Nielsen + Shneiderman."),
        ("5. Evaluación heurística (4 min)", "Top 3 graves: control (1→3), flexibilidad (1→3), prevención errores (parcial)."),
        ("6. Reflexión (3 min)", "Documentos como evidencia; Xojo futuro; prueba con 3 usuarios."),
        ("7. Conclusión (1 min)", "«El diseño mejoró cuando el biólogo dejó de pelear con la terminal y volvió a mirar el arrecife.»"),
    ]
    for t, c in blocks:
        doc.add_heading(t, level=1)
        doc.add_paragraph(c)
    doc.add_heading("Diapositivas sugeridas", level=1)
    slides = [
        "Portada + matrícula",
        "Problema + foto de arrecife",
        "Perfil biólogo (tabla Informe_Biologo)",
        "Línea tiempo: Eval 1.6 → cambios → Eval 2.7",
        "Captura Cargar (drag-and-drop)",
        "Captura Revisión (puntos verde/rojo)",
        "Tabla 3 hallazgos graves",
        "Atajos J K R D (ATAJOS_TECLADO)",
        "Valoraciones biólogo 4–5/5",
        "Trabajo futuro Xojo",
        "Cierre",
    ]
    for s in slides:
        doc.add_paragraph(s, style="List Bullet")
    doc.save(PRESENTATION)
    print(f"Guion: {PRESENTATION}")


if __name__ == "__main__":
    build_report()
    build_presentation()
